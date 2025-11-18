"""
PDF Comparison Tool - Identifies differences between two PDFs
Handles text, tables, and moved content across pages
Outputs: PDF, DOCX, or Excel
"""

import os
import sys
from pathlib import Path
from typing import List, Dict, Tuple, Set
from difflib import SequenceMatcher, unified_diff
from collections import defaultdict
import json
import re

# Required packages (install with: pip install -r requirements.txt)
# pdfplumber, python-docx, openpyxl, pillow, reportlab, fuzzywuzzy, python-Levenshtein

try:
    import pdfplumber
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from fuzzywuzzy import fuzz
    from PIL import Image
    import io
except ImportError as e:
    print(f"❌ Missing required package: {e}")
    print("\n📦 Install required packages with:")
    print("pip install pdfplumber python-docx openpyxl pillow reportlab fuzzywuzzy python-Levenshtein")
    sys.exit(1)


class PDFComparisonTool:
    def __init__(self, pdf1_path: str, pdf2_path: str, similarity_threshold: int = 85):
        """
        Initialize PDF comparison tool
        
        Args:
            pdf1_path: Path to first PDF (old version)
            pdf2_path: Path to second PDF (new version)
            similarity_threshold: Fuzzy matching threshold (0-100)
        """
        self.pdf1_path = pdf1_path
        self.pdf2_path = pdf2_path
        self.similarity_threshold = similarity_threshold
        
        # Validation
        if not Path(pdf1_path).exists():
            raise FileNotFoundError(f"❌ PDF 1 not found: {pdf1_path}")
        if not Path(pdf2_path).exists():
            raise FileNotFoundError(f"❌ PDF 2 not found: {pdf2_path}")
        
        self.pdf1_content = []
        self.pdf2_content = []
        self.differences = {
            'added': [],
            'removed': [],
            'modified': [],
            'moved': []
        }
        
    def extract_content(self, pdf_path: str) -> List[Dict]:
        """
        Extract text and tables from PDF with metadata
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of content dictionaries with page info and text/tables
        """
        print(f"📄 Extracting content from: {Path(pdf_path).name}")
        content = []
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"   Total pages: {total_pages}")
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_content = {
                        'page': page_num,
                        'text': [],
                        'tables': [],
                        'raw_text': ''
                    }
                    
                    # Extract text
                    text = page.extract_text() or ""
                    if text.strip():
                        page_content['raw_text'] = text
                        # Split into paragraphs (by double newline or significant text blocks)
                        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
                        page_content['text'] = paragraphs
                    
                    # Extract tables
                    try:
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables):
                                table_dict = {
                                    'index': table_idx,
                                    'data': table,
                                    'text_representation': self._table_to_text(table)
                                }
                                page_content['tables'].append(table_dict)
                    except Exception as e:
                        print(f"   ⚠️  Warning extracting table from page {page_num}: {str(e)}")
                    
                    content.append(page_content)
                    
                print(f"   ✓ Extraction complete")
        except Exception as e:
            print(f"❌ Error extracting PDF: {e}")
            raise
        
        return content
    
    @staticmethod
    def _table_to_text(table: List[List]) -> str:
        """Convert table to text representation for comparison"""
        rows = []
        for row in table:
            row_text = " | ".join([str(cell) if cell else "" for cell in row])
            rows.append(row_text)
        return "\n".join(rows)
    
    def normalize_text(self, text: str) -> str:
        """Normalize text for better comparison"""
        # Remove extra whitespace
        text = " ".join(text.split())
        # Convert to lowercase for comparison
        text = text.lower()
        # Remove special characters but keep alphanumeric
        text = re.sub(r'[^\w\s]', '', text)
        return text
    
    def find_similar_content(self, content_list: List[str], target: str, 
                            threshold: int = None) -> Tuple[bool, int, int]:
        """
        Find similar content across PDFs using fuzzy matching
        Handles case where content moved between pages
        
        Args:
            content_list: List of content to search in
            target: Content to find
            threshold: Similarity threshold
            
        Returns:
            Tuple of (found: bool, similarity_score: int, index: int)
        """
        if threshold is None:
            threshold = self.similarity_threshold
        
        best_match = 0
        best_index = -1
        
        normalized_target = self.normalize_text(target)
        
        for idx, content in enumerate(content_list):
            normalized_content = self.normalize_text(content)
            
            # Use token_set_ratio for better matching of moved/modified content
            similarity = fuzz.token_set_ratio(normalized_target, normalized_content)
            
            if similarity > best_match:
                best_match = similarity
                best_index = idx
        
        found = best_match >= threshold
        return found, best_match, best_index
    
    def extract_all_text_items(self, content_list: List[Dict]) -> Tuple[Dict, Set]:
        """
        Extract all unique text items from content
        Returns: (mapping of normalized text to original, set of normalized texts)
        """
        text_mapping = {}
        text_set = set()
        
        for page_content in content_list:
            # Add paragraph text
            for para in page_content['text']:
                normalized = self.normalize_text(para)
                if normalized:  # Only add non-empty
                    text_mapping[normalized] = {
                        'original': para,
                        'page': page_content['page'],
                        'type': 'text'
                    }
                    text_set.add(normalized)
            
            # Add table representations
            for table in page_content['tables']:
                normalized = self.normalize_text(table['text_representation'])
                if normalized:
                    text_mapping[normalized] = {
                        'original': table['text_representation'][:100] + '...',
                        'page': page_content['page'],
                        'type': 'table'
                    }
                    text_set.add(normalized)
        
        return text_mapping, text_set
    
    def compare_pdfs(self):
        """Main comparison logic"""
        print("\n🔍 Starting PDF Comparison...\n")
        
        # Extract content from both PDFs
        self.pdf1_content = self.extract_content(self.pdf1_path)
        self.pdf2_content = self.extract_content(self.pdf2_path)
        
        # Get all unique text items from each PDF
        pdf1_mapping, pdf1_texts = self.extract_all_text_items(self.pdf1_content)
        pdf2_mapping, pdf2_texts = self.extract_all_text_items(self.pdf2_content)
        
        print("\n📊 Analyzing Differences...\n")
        
        # Find added content (in PDF2 but not in PDF1)
        for normalized_text in pdf2_texts:
            found, score, _ = self.find_similar_content(list(pdf1_texts), normalized_text)
            
            if not found:
                original_text = pdf2_mapping[normalized_text]
                self.differences['added'].append({
                    'content': original_text['original'],
                    'page': original_text['page'],
                    'type': original_text['type'],
                    'similarity': 0
                })
                print(f"➕ ADDED (Page {original_text['page']}): {original_text['original'][:60]}...")
        
        # Find removed content (in PDF1 but not in PDF2)
        for normalized_text in pdf1_texts:
            found, score, _ = self.find_similar_content(list(pdf2_texts), normalized_text)
            
            if not found:
                original_text = pdf1_mapping[normalized_text]
                self.differences['removed'].append({
                    'content': original_text['original'],
                    'page': original_text['page'],
                    'type': original_text['type'],
                    'similarity': 0
                })
                print(f"➖ REMOVED (Page {original_text['page']}): {original_text['original'][:60]}...")
        
        # Find moved content (same content but different pages)
        for normalized_text in pdf1_texts:
            if normalized_text in pdf2_texts:
                pdf1_page = pdf1_mapping[normalized_text]['page']
                pdf2_page = pdf2_mapping[normalized_text]['page']
                
                if pdf1_page != pdf2_page:
                    self.differences['moved'].append({
                        'content': pdf1_mapping[normalized_text]['original'],
                        'from_page': pdf1_page,
                        'to_page': pdf2_page,
                        'type': pdf1_mapping[normalized_text]['type']
                    })
                    print(f"🔄 MOVED: Page {pdf1_page} → Page {pdf2_page}: {pdf1_mapping[normalized_text]['original'][:60]}...")
        
        print(f"\n✅ Comparison Complete!")
        print(f"   Added: {len(self.differences['added'])}")
        print(f"   Removed: {len(self.differences['removed'])}")
        print(f"   Moved: {len(self.differences['moved'])}")
        print(f"   Modified: {len(self.differences['modified'])}")
    
    def generate_docx_report(self, output_path: str):
        """Generate DOCX report"""
        print(f"\n📝 Generating DOCX report: {output_path}")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('PDF Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_heading('Summary', level=1)
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Light Grid Accent 1'
        
        summary_data = [
            ('Total Added Items', str(len(self.differences['added']))),
            ('Total Removed Items', str(len(self.differences['removed']))),
            ('Total Moved Items', str(len(self.differences['moved']))),
            ('Total Modified Items', str(len(self.differences['modified']))),
        ]
        
        for row_idx, (label, value) in enumerate(summary_data):
            cells = summary_table.rows[row_idx + 1].cells
            cells[0].text = label
            cells[1].text = value
        
        # Added Content
        if self.differences['added']:
            doc.add_heading('Added Content', level=1)
            for idx, item in enumerate(self.differences['added'], 1):
                p = doc.add_paragraph(f"Item {idx} (Page {item['page']}, Type: {item['type']})", style='Heading 3')
                p.runs[0].font.color.rgb = RGBColor(0, 176, 0)  # Green
                doc.add_paragraph(item['content'][:500])
        
        # Removed Content
        if self.differences['removed']:
            doc.add_heading('Removed Content', level=1)
            for idx, item in enumerate(self.differences['removed'], 1):
                p = doc.add_paragraph(f"Item {idx} (Page {item['page']}, Type: {item['type']})", style='Heading 3')
                p.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
                doc.add_paragraph(item['content'][:500])
        
        # Moved Content
        if self.differences['moved']:
            doc.add_heading('Moved Content', level=1)
            for idx, item in enumerate(self.differences['moved'], 1):
                p = doc.add_paragraph(f"Item {idx} (Page {item['from_page']} → {item['to_page']})", style='Heading 3')
                p.runs[0].font.color.rgb = RGBColor(255, 153, 0)  # Orange
                doc.add_paragraph(item['content'][:500])
        
        doc.save(output_path)
        print(f"✅ DOCX report saved: {output_path}")
    
    def generate_excel_report(self, output_path: str):
        """Generate Excel report"""
        print(f"\n📊 Generating Excel report: {output_path}")
        
        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet
        
        # Define styles
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        
        added_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
        removed_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
        moved_fill = PatternFill(start_color="FFE699", end_color="FFE699", fill_type="solid")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # Summary Sheet
        ws_summary = wb.create_sheet('Summary')
        ws_summary['A1'] = 'PDF Comparison Report - Summary'
        ws_summary['A1'].font = Font(bold=True, size=14)
        
        summary_data = [
            ['Metric', 'Count'],
            ['Added Items', len(self.differences['added'])],
            ['Removed Items', len(self.differences['removed'])],
            ['Moved Items', len(self.differences['moved'])],
            ['Total Differences', sum(len(v) for v in self.differences.values())]
        ]
        
        for row_idx, row_data in enumerate(summary_data, 1):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
                if row_idx == 1:
                    cell.fill = header_fill
                    cell.font = header_font
                cell.border = border
        
        ws_summary.column_dimensions['A'].width = 25
        ws_summary.column_dimensions['B'].width = 15
        
        # Added Content Sheet
        if self.differences['added']:
            ws_added = wb.create_sheet('Added')
            headers = ['#', 'Page', 'Type', 'Content']
            for col_idx, header in enumerate(headers, 1):
                cell = ws_added.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
            
            for row_idx, item in enumerate(self.differences['added'], 2):
                ws_added.cell(row=row_idx, column=1, value=row_idx - 1).fill = added_fill
                ws_added.cell(row=row_idx, column=2, value=item['page']).fill = added_fill
                ws_added.cell(row=row_idx, column=3, value=item['type']).fill = added_fill
                ws_added.cell(row=row_idx, column=4, value=str(item['content'])[:100]).fill = added_fill
            
            ws_added.column_dimensions['A'].width = 5
            ws_added.column_dimensions['B'].width = 10
            ws_added.column_dimensions['C'].width = 12
            ws_added.column_dimensions['D'].width = 60
        
        # Removed Content Sheet
        if self.differences['removed']:
            ws_removed = wb.create_sheet('Removed')
            headers = ['#', 'Page', 'Type', 'Content']
            for col_idx, header in enumerate(headers, 1):
                cell = ws_removed.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
            
            for row_idx, item in enumerate(self.differences['removed'], 2):
                ws_removed.cell(row=row_idx, column=1, value=row_idx - 1).fill = removed_fill
                ws_removed.cell(row=row_idx, column=2, value=item['page']).fill = removed_fill
                ws_removed.cell(row=row_idx, column=3, value=item['type']).fill = removed_fill
                ws_removed.cell(row=row_idx, column=4, value=str(item['content'])[:100]).fill = removed_fill
            
            ws_removed.column_dimensions['A'].width = 5
            ws_removed.column_dimensions['B'].width = 10
            ws_removed.column_dimensions['C'].width = 12
            ws_removed.column_dimensions['D'].width = 60
        
        # Moved Content Sheet
        if self.differences['moved']:
            ws_moved = wb.create_sheet('Moved')
            headers = ['#', 'From Page', 'To Page', 'Type', 'Content']
            for col_idx, header in enumerate(headers, 1):
                cell = ws_moved.cell(row=1, column=col_idx, value=header)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
            
            for row_idx, item in enumerate(self.differences['moved'], 2):
                ws_moved.cell(row=row_idx, column=1, value=row_idx - 1).fill = moved_fill
                ws_moved.cell(row=row_idx, column=2, value=item['from_page']).fill = moved_fill
                ws_moved.cell(row=row_idx, column=3, value=item['to_page']).fill = moved_fill
                ws_moved.cell(row=row_idx, column=4, value=item['type']).fill = moved_fill
                ws_moved.cell(row=row_idx, column=5, value=str(item['content'])[:100]).fill = moved_fill
            
            ws_moved.column_dimensions['A'].width = 5
            ws_moved.column_dimensions['B'].width = 12
            ws_moved.column_dimensions['C'].width = 10
            ws_moved.column_dimensions['D'].width = 12
            ws_moved.column_dimensions['E'].width = 50
        
        wb.save(output_path)
        print(f"✅ Excel report saved: {output_path}")
    
    def generate_pdf_report(self, output_path: str):
        """Generate PDF report using ReportLab"""
        print(f"\n📄 Generating PDF report: {output_path}")
        
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak, Heading1, Heading2
            from reportlab.enum.text import TA_CENTER, TA_LEFT
        except ImportError:
            print("⚠️  ReportLab not installed. Install with: pip install reportlab")
            print("📝 Falling back to DOCX format...")
            self.generate_docx_report(output_path.replace('.pdf', '.docx'))
            return
        
        # Create PDF
        doc = SimpleDocTemplate(output_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#4472C4'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph('PDF Comparison Report', title_style))
        story.append(Spacer(1, 12))
        
        # Summary Table
        summary_data = [
            ['Metric', 'Count'],
            ['Added Items', str(len(self.differences['added']))],
            ['Removed Items', str(len(self.differences['removed']))],
            ['Moved Items', str(len(self.differences['moved']))],
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 2*inch])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#4472C4')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(summary_table)
        story.append(Spacer(1, 20))
        
        # Added Content Section
        if self.differences['added']:
            story.append(Paragraph('Added Content', styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for idx, item in enumerate(self.differences['added'][:10], 1):  # Limit to first 10
                content_text = str(item['content'])[:150]
                story.append(Paragraph(f"<b>Item {idx}</b> (Page {item['page']}): {content_text}...", styles['Normal']))
                story.append(Spacer(1, 6))
            
            story.append(Spacer(1, 12))
        
        # Removed Content Section
        if self.differences['removed']:
            story.append(Paragraph('Removed Content', styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for idx, item in enumerate(self.differences['removed'][:10], 1):  # Limit to first 10
                content_text = str(item['content'])[:150]
                story.append(Paragraph(f"<b>Item {idx}</b> (Page {item['page']}): {content_text}...", styles['Normal']))
                story.append(Spacer(1, 6))
            
            story.append(Spacer(1, 12))
        
        # Moved Content Section
        if self.differences['moved']:
            story.append(Paragraph('Moved Content', styles['Heading2']))
            story.append(Spacer(1, 12))
            
            for idx, item in enumerate(self.differences['moved'][:10], 1):  # Limit to first 10
                content_text = str(item['content'])[:150]
                story.append(Paragraph(f"<b>Item {idx}</b> (Page {item['from_page']} → {item['to_page']}): {content_text}...", styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        print(f"✅ PDF report saved: {output_path}")


def main():
    """Main function"""
    print("=" * 70)
    print("🔍 PDF COMPARISON TOOL - Compare Two PDFs & Generate Reports 🔍")
    print("=" * 70)
    print()
    
    # Get user input
    pdf1 = input("📄 Enter path to PDF 1 (Old Version): ").strip()
    pdf2 = input("📄 Enter path to PDF 2 (New Version): ").strip()
    
    if not pdf1 or not pdf2:
        print("❌ PDF paths cannot be empty!")
        return
    
    print("\n📋 Output Format Options:")
    print("   1. Excel (.xlsx)")
    print("   2. Word (.docx)")
    print("   3. PDF (.pdf)")
    print("   4. All formats")
    
    choice = input("\n🎯 Select output format (1-4): ").strip()
    
    try:
        # Initialize comparison tool
        tool = PDFComparisonTool(pdf1, pdf2, similarity_threshold=85)
        
        # Perform comparison
        tool.compare_pdfs()
        
        # Generate reports based on choice
        output_dir = Path("comparison_reports")
        output_dir.mkdir(exist_ok=True)
        
        timestamp = __import__('datetime').datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if choice in ['1', '4']:
            excel_output = output_dir / f"comparison_report_{timestamp}.xlsx"
            tool.generate_excel_report(str(excel_output))
        
        if choice in ['2', '4']:
            docx_output = output_dir / f"comparison_report_{timestamp}.docx"
            tool.generate_docx_report(str(docx_output))
        
        if choice in ['3', '4']:
            pdf_output = output_dir / f"comparison_report_{timestamp}.pdf"
            tool.generate_pdf_report(str(pdf_output))
        
        print("\n" + "=" * 70)
        print("✅ PDF COMPARISON COMPLETE!")
        print(f"📁 Reports saved in: {output_dir.absolute()}")
        print("=" * 70)
        
    except Exception as e:
        print(f"\n❌ Error occurred: {e}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    main()
